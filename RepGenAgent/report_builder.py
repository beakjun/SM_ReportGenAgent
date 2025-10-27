import os
import json
import io
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement

import warnings
warnings.filterwarnings('ignore')

# ============================================================
# 1. JSON 데이터 전처리
# ============================================================
def build_doc_structure(loaded_dict: dict) -> dict:
    """
    loaded_dict에서 문서 구조를 추출해 계층형 JSON 구조로 반환하는 함수.
    
    Parameters
    ----------
    loaded_dict : dict
        'documents', 'results', 'report_title', 'report_conclusion' 등을 포함한 원본 딕셔너리

    Returns
    -------
    dict
        {
            "doc_title": str,
            "report_conclusion": str,
            "blocks": list
        }
    """

    # 필요한 key만 안전하게 추출
    doc_text = loaded_dict.get('documents', '')
    results = loaded_dict.get('results', [])
    doc_title = loaded_dict.get('report_title', '')
    report_conclusion = loaded_dict.get('report_conclusion', '')

    # "주요 목차" 구간만 추출
    match = re.search(r'주요 목차:(.*?)(?=-목차 콘텐츠:)', doc_text, re.DOTALL)
    section = match.group(1).strip() if match else ""

    # 주요 목차 리스트 생성
    headings = re.findall(r'\d+\.\s*[^0-9\n]+', section)
    headings = [h.strip() for h in headings]

    # 블록 구성
    blocks = []
    for h in headings:
        prefix = h.split('.')[0] + '.'

        # 하위 블록 (예: 3.1., 3.2.)
        children = [r for r in results if r['contents_id'].startswith(prefix) and r['contents_id'] != h]
        # 상위 블록
        parent_block = next((r for r in results if r['contents_id'] == h), None)

        if children:
            child_type = children[0].get('type', '')
            merged_type = f"Multi{child_type}" if child_type else "Multi"

            block = {
                "contents_id": h,
                "type": merged_type,
                "contents": parent_block.get('contents') if parent_block else '',
                "results": children
            }
            blocks.append(block)
        elif parent_block:
            blocks.append(parent_block)

    # 최종 JSON 구조
    output = {
        "doc_title": doc_title,
        "report_conclusion": report_conclusion,
        "blocks": blocks
    }

    return output


# ============================================================
# 2. 문서 스타일 설정
# ============================================================
def set_default_style(doc):
    """기본 폰트와 페이지 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT

    margin = Cm(1.27)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = margin
    section.footer_distance = margin


def add_block_heading(target, text, font_name="Malgun Gothic", font_size=14):
    """
    블록 제목 생성
    - target: Document 또는 Cell
    - text: 제목 텍스트
    - font_name: 한글 폰트
    - font_size: pt
    """
    # container가 list/tuple로 넘어오면 첫 셀 사용
    if isinstance(target, (list, tuple)):
        target = target[0]

    para = target.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    try:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass

    # Heading 느낌 여백
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)

    return para

# ============================================================
# 3. 레이아웃 템플릿
# ============================================================
def set_table_autofit(table, autofit=True):
    """
    표를 Word의 '창에 자동 맞춤' 모드처럼 설정하는 함수
    """
    tbl = table._element
    tblPr = tbl.tblPr

    # 기존 tblLayout 제거
    for e in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(e)

    # <w:tblW> 노드가 없으면 생성
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)

    if autofit:
        tblW.set(qn('w:type'), 'auto')   # 창에 맞춤
    else:
        tblW.set(qn('w:type'), 'dxa')    # 고정 폭
        tblW.set(qn('w:w'), '9000')      # 적당한 크기 (단위: 1/20 pt)

def create_report_template(doc, layout_type="1col"):
    """
    보고서 레이아웃 템플릿
    layout_type:
        - "1col" : 본문 전체 폭
        - "2col" : 두 개의 컬럼
        - "full_table" : 표/차트 전체 폭
    """
    if layout_type == "1col":
        table = doc.add_table(rows=1, cols=1)
        set_table_autofit(table, autofit=True)
        return table.cell(0, 0)

    elif layout_type == "2col":
        table = doc.add_table(rows=1, cols=2)
        set_table_autofit(table, autofit=True)
        return table.cells  # (왼쪽셀, 오른쪽셀)

    elif layout_type == "full_table":
        table = doc.add_table(rows=1, cols=1)
        set_table_autofit(table, autofit=True)
        return table.cell(0, 0)

    else:
        raise ValueError("지원하지 않는 layout_type 입니다.")

# ============================================================
# 4. 기본 테이블 / 셀 렌더링
# ============================================================
def create_table(doc, n_rows, n_cols, style=None):
    """문서에 n_rows x n_cols 테이블 추가 후 반환"""
    table = doc.add_table(rows=n_rows, cols=n_cols)
    if style:
        table.style = style
    return table

def parse_table_data(value):
    """문자열 → list[dict] 변환 유틸 (깨진 JSON도 복구 시도)"""
    if isinstance(value, (list, dict)):
        return value

    if not isinstance(value, str):
        return None

    # JSON 시도
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        pass

    # Python literal 시도
    try:
        return ast.literal_eval(value)
    except Exception:
        pass

    # 따옴표 교정 후 재시도
    try:
        repaired = re.sub(r"'", '"', value)
        return json.loads(repaired)
    except Exception:
        return None


def render_table(values, container, style="Table Grid"):
    """
    리스트[딕셔너리] 기반 표를 Document 또는 Cell에 생성
    - 헤더: 회색 배경 + 굵게 + 가운데 정렬
    - 데이터: 가운데 정렬
    """
    values = parse_table_data(values)

    if not values:
        # 완전히 파싱 실패 시 그냥 텍스트로 출력
        return container.add_paragraph(str(values))

    # dict 단일 → list로 변환
    if isinstance(values, dict):
        values = [values]

    # 데이터가 list[dict] 형태인지 확인
    if not isinstance(values, list) or not isinstance(values[0], dict):
        return container.add_paragraph(str(values))

    headers = list(values[0].keys())

    table = container.add_table(rows=1, cols=len(headers))
    table.style = style

    # === 헤더 ===
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i]._tc.get_or_add_tcPr().append(
            parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
        )
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # === 데이터 ===
    for row in values:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            row_cells[i].text = str(row.get(h, ""))
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table



# ============================================================
# 5. 텍스트 / 표 / 차트 블록
# ============================================================
def add_text_section(doc, block, container=None):
    """텍스트 블록 처리 (레이아웃 컨테이너 안에 넣을 수 있음)"""
    target = container if container else doc

    # 제목 처리
    add_block_heading(target, f"{block['contents_id']}")

    paragraph = target.add_paragraph()
    text = block["results"]
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True
            
def flatten_table_string(data_str):
    """
    문자열 형태의 테이블 데이터를 단일 row dict로 변환
    - {0: value} → value
    - [value] → value
    - 이미 단일 값이면 그대로
    """
    import ast

    # 문자열 → dict
    try:
        data_dict = ast.literal_eval(data_str)
    except Exception:
        return {}

    flat_dict = {}
    for k, v in data_dict.items():
        if isinstance(v, dict) and 0 in v:
            flat_dict[k] = v[0]
        elif isinstance(v, list) and len(v) == 1:
            flat_dict[k] = v[0]
        else:
            flat_dict[k] = v
    return flat_dict

def add_table_section(doc, block, container=None):
    """Table 블록 처리"""
    target = container if container else doc

    # 제목 처리
    add_block_heading(target, f"{block['contents_id']}")

    # flatten_table_string() 적용
    try:
        table_data = flatten_table_string(block["results"])
    except Exception:
        table_data = block["results"]

    render_table(table_data, target)


def add_chart_section(doc, block, container=None):
    """Chart 블록 처리"""
    target = container if container else doc
 
    # 제목 처리
    add_block_heading(target, f"{block['contents_id']}")

    img_stream = render_chart_to_stream(block["results"])
    paragraph = target.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(img_stream, width=Inches(6.0))  # height=Inches(3.5)


# ============================================================
# 6. 멀티테이블 / 멀티차트 블록
# ============================================================
def multi_table_text(doc, block):
    """MultiTable 블록 처리"""
    add_block_heading(doc, f"{block['contents_id']}")

    child_blocks = block["results"]
    n_cols = len(child_blocks)
    table = create_table(doc, 2, n_cols, style="Table Grid")

    for idx, child in enumerate(child_blocks):
        # 1행: 소제목
        header_cell = table.cell(0, idx)
        header_cell.text = child["contents_id"]
        header_cell._tc.get_or_add_tcPr().append(
            parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
        )
        for p in header_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True

        # 2행: 데이터 요약
        data_cell = table.cell(1, idx)
        values = flatten_table_string(child["results"])  # 단일 row 전처리

        # Rank 정보 매핑 (대소문자 무시)
        rank_map = {
            k.lower().replace("_rank", ""): (v[0] if isinstance(v, list) else v)
            for k, v in values.items()
            if k.lower().endswith("_rank")
        }

        lines = []
        for k, v in values.items():
            # Rank 컬럼 자체는 건너뛰기
            if k.lower().endswith("_rank"):
                continue

            val = v[0] if isinstance(v, list) else v
            rank_val = rank_map.get(k.lower())

            if rank_val is not None:
                lines.append(f"{k}: {val} ({int(rank_val)}위)")
            else:
                lines.append(f"{k}: {val}")

        data_cell.text = "\n".join(lines)

    return table


def multi_table_chart(doc, block):
    """MultiChart 블록 처리"""
    # 제목 처리
    add_block_heading(doc, f"{block['contents_id']}")
    
    child_blocks = block["results"]
    n_cols = len(child_blocks)
    table = create_table(doc, 3, n_cols)

    # 1행: 소제목
    for j, child in enumerate(child_blocks):
        cell = table.cell(0, j)
        cell.text = child["contents_id"]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(9)
        cell._tc.get_or_add_tcPr().append(
            parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
        )

    # 2행: 차트
    for j, child in enumerate(child_blocks):
        cell = table.cell(1, j)
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_stream = render_chart_to_stream(child["results"])
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(2.5))

    # 3행: 데이터 테이블
    for j, child in enumerate(child_blocks):
        cell = table.cell(2, j)
        render_table(child["results"], cell)

    return table


# ============================================================
# 7. 차트 렌더링
# ============================================================
def render_chart_to_stream(code_value):
    """코드 실행 후 차트를 PNG 메모리 스트림으로 반환"""
    exec(code_value, globals())
    fig = plt.gcf()
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format="png", bbox_inches="tight")
    img_stream.seek(0)
    plt.close(fig)
    return img_stream


# ============================================================
# 8. 보고서 빌드
# ============================================================
def build_report(data_dict, output_file="report.docx", output_dir=None):
    """
    JSON 데이터를 기반으로 보고서 작성
    
    Parameters
    ----------
    data_dict : dict
        build_doc_structure로 생성된 문서 구조
    output_file : str
        저장할 파일명 (예: report.docx)
    output_dir : str or None
        저장할 폴더 경로. None이면 현재 작업 디렉토리에 저장
    """
    # 저장 경로 처리
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_file)
    else:
        output_path = output_file

    doc = Document()
    set_default_style(doc)

    # === 문서 제목 ===
    heading = doc.add_heading(data_dict["doc_title"], level=0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)

    # === 문서 요약 ===
    intro = data_dict.get("report_conclusion", "")
    if intro:
        #add_block_heading(doc, "0. 요약")   # 제목 수동 추가
        doc.add_paragraph("")  # 빈 단락으로 줄띄움

        # **굵게** 처리 로직 적용
        paragraph = doc.add_paragraph()
        parts = intro.split("**")
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True

    # === 본문 블록 ===
    for block in data_dict["blocks"]:
        btype = block["type"].lower()

        # === 레이아웃 선택 ===
        if btype in ("text", "table"):
            container = create_report_template(doc, "1col")
        elif btype == "chart":
            container = create_report_template(doc, "full_table")
        else:
            container = None  # 멀티형

        # === 블록 렌더링 ===
        if btype == "text":
            add_text_section(doc, block, container)
        elif btype == "table":
            add_table_section(doc, block, container)
        elif btype == "chart":
            add_chart_section(doc, block, container)
        elif btype == "multitable":
            multi_table_text(doc, block)
        elif btype == "multichart":
            multi_table_chart(doc, block)

    doc.save(output_path)
    print(f"보고서가 생성되었습니다: {output_path}")